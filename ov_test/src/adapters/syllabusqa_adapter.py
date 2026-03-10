# src/adapters/syllabusqa_adapter.py
"""
SyllabusQA 数据集适配器

SyllabusQA 是一个课程大纲问答数据集，包含 39 个课程大纲和 5078 个问题。
每个问题针对特定的课程大纲（syllabus）提问，答案类型包括：
- single factual: 单一事实问题
- multi factual: 多事实问题
- single reasoning: 单一推理问题
- multi reasoning: 多推理问题
- summarization: 总结类问题
- yes/no: 是/否问题
- no answer: 无法从大纲中找到答案

数据集特点：
1. 每个问题只有一个答案
2. 每个答案有对应的 answer_span（证据片段）
3. reasoning 类型问题有推理步骤
4. 原始文档为 docx 格式

适配器功能：
- data_prepare: 将 docx 转换为 Markdown 格式
- load_and_transform: 解析 QA 数据，category 存储问题类型
- build_prompt: 构建问答提示词
- post_process_answer: 后处理 LLM 输出
"""

import json
import os
import csv
from typing import List, Dict, Any

from .base import BaseAdapter, StandardDoc, StandardSample, StandardQA

# 问答提示词模板
QA_PROMPT = """Based on the above context, write an answer in the form of a short phrase for the following question. Answer with exact words from the context whenever possible.

Question: {} Short answer:
"""

# 无法回答时的规则说明
MISSING_RULE = "If no information is available to answer the question, write 'Not mentioned'."


class SyllabusQAAdapter(BaseAdapter):
    """
    专门用于处理 SyllabusQA 数据集的适配器。
    
    将课程大纲（docx）转换为 Markdown 文档，
    并将 QA 数据转换为标准化的 StandardSample 格式。
    
    Attributes:
        raw_file_path: 原始 CSV 数据文件路径
        syllabus_dir: docx 文件目录路径
        logger: 日志记录器
    """
    
    def __init__(self, raw_file_path: str, **kwargs):
        """
        初始化 SyllabusQAAdapter。
        
        Args:
            raw_file_path: 原始数据文件路径（CSV 或合并后的 JSON）
            **kwargs: 其他参数（被忽略，用于兼容性）
        """
        super().__init__(raw_file_path)
        # docx 文件目录，默认为数据目录下的 syllabi 子目录
        if os.path.isdir(raw_file_path):
            self.syllabus_dir = os.path.join(raw_file_path, 'syllabi')
        else:
            self.syllabus_dir = os.path.join(os.path.dirname(raw_file_path), 'syllabi')
    
    def data_prepare(self, doc_dir: str) -> List[StandardDoc]:
        """
        加载原始 docx 文件并转换为 OpenViking 友好格式。
        
        只处理 CSV 中涉及的 syllabus 文档，避免处理无关文档。
        需要安装 python-docx 库来解析 docx 文件。
        
        Args:
            doc_dir: 文档输出目录路径
            
        Returns:
            List[StandardDoc]: 标准化文档对象列表
            
        Raises:
            FileNotFoundError: syllabus 目录不存在
        """
        if not os.path.exists(self.syllabus_dir):
            raise FileNotFoundError(f"Syllabus directory not found: {self.syllabus_dir}")

        res: List[StandardDoc] = []
        os.makedirs(doc_dir, exist_ok=True)
        
        # 获取 CSV 中涉及的 syllabus_name 列表
        required_syllabi = self._get_required_syllabi()
        self.logger.info(f"[SyllabusQAAdapter] Required syllabi from CSV: {len(required_syllabi)}")
        
        # 获取所有 docx 文件
        docx_files = [f for f in os.listdir(self.syllabus_dir) if f.endswith('.docx')]
        
        for docx_file in docx_files:
            syllabus_id = docx_file.replace('.docx', '')
            
            # 只处理 CSV 中涉及的 syllabus
            if syllabus_id not in required_syllabi:
                continue
            
            docx_path = os.path.join(self.syllabus_dir, docx_file)
            
            try:
                # 转换 docx 为 Markdown
                doc_content = self._convert_docx_to_markdown(docx_path)
                
                doc_path = os.path.join(doc_dir, f"{syllabus_id}_doc.md")
                with open(doc_path, "w", encoding="utf-8") as f:
                    f.write(doc_content)
                res.append(StandardDoc(syllabus_id, [doc_path]))
            except Exception as e:
                self.logger.error(f"[syllabusqa adapter] doc:{syllabus_id} prepare error {e}")
                # 如果 python-docx 未安装，尝试使用纯文本方式
                if "No module named 'docx'" in str(e):
                    self.logger.warning("python-docx not installed, skipping docx conversion")
                    break
                raise e
        
        self.logger.info(f"[SyllabusQAAdapter] Processed {len(res)} syllabus documents")
        return res

    def _get_required_syllabi(self) -> set:
        """
        获取 CSV 中涉及的 syllabus_name 列表。
        
        Returns:
            set: syllabus_name 集合
        """
        required = set()
        
        # 判断数据源类型
        if self.raw_file_path.endswith('.csv'):
            csv_files = [self.raw_file_path]
        elif os.path.isdir(self.raw_file_path):
            csv_files = [os.path.join(self.raw_file_path, f) 
                        for f in os.listdir(self.raw_file_path) 
                        if f.endswith('.csv')]
        else:
            return required
        
        for csv_file in csv_files:
            if not os.path.exists(csv_file):
                continue
            with open(csv_file, 'r', encoding='utf-8') as f:
                import csv
                reader = csv.DictReader(f)
                for row in reader:
                    syllabus_name = row.get('syllabus_name', '')
                    if syllabus_name:
                        required.add(syllabus_name)
        
        return required

    def _convert_docx_to_markdown(self, docx_path: str) -> str:
        """
        将 docx 文件转换为 Markdown 字符串。
        
        Args:
            docx_path: docx 文件路径
            
        Returns:
            str: Markdown 格式的内容
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        doc = Document(docx_path)
        md_lines = []
        
        # 提取文件名作为标题
        filename = os.path.basename(docx_path).replace('.docx', '')
        md_lines.append(f"# {filename}")
        md_lines.append("")
        
        # 遍历所有段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 检查是否为标题样式
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        md_lines.append(f"{'#' * level_num} {text}")
                    except ValueError:
                        md_lines.append(f"## {text}")
                else:
                    md_lines.append(text)
                md_lines.append("")
        
        # 提取表格
        for table in doc.tables:
            md_lines.append("## Table")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
            md_lines.append("")
        
        return "\n".join(md_lines)

    def load_and_transform(self) -> List[StandardSample]:
        """
        加载原始 CSV 数据并转换为标准化的 StandardSample 对象列表。
        
        处理逻辑：
        1. 读取 CSV 文件（支持单个 CSV 或目录）
        2. 按 syllabus_name 分组
        3. 问题格式化为 "Based on the syllabus "{syllabus_name}", {question}"
        4. category 存储 question_type
        5. answer_span 作为 evidence
        
        Returns:
            List[StandardSample]: 标准化样本对象列表
            
        Raises:
            FileNotFoundError: 原始数据文件不存在
        """
        # 判断是 CSV 文件还是 JSON 文件
        if self.raw_file_path.endswith('.json'):
            return self._load_from_json()
        elif self.raw_file_path.endswith('.csv'):
            return self._load_from_csv([self.raw_file_path])
        elif os.path.isdir(self.raw_file_path):
            # 目录，查找所有 CSV 文件
            csv_files = [os.path.join(self.raw_file_path, f) 
                        for f in os.listdir(self.raw_file_path) 
                        if f.endswith('.csv')]
            return self._load_from_csv(csv_files)
        else:
            raise FileNotFoundError(f"Unsupported file format: {self.raw_file_path}")

    def _load_from_json(self) -> List[StandardSample]:
        """
        从 JSON 文件加载数据。
        """
        if not os.path.exists(self.raw_file_path):
            raise FileNotFoundError(f"Raw data file not found: {self.raw_file_path}")

        with open(self.raw_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        standard_samples = []

        for syllabus_name, qa_list in data.items():
            qa_pairs = []
            
            for qa_item in qa_list:
                question = qa_item.get("question", "")
                answer = qa_item.get("answer", "")
                question_type = qa_item.get("question_type", "")
                qa_id = qa_item.get("id", "")
                
                # 跳过 "no answer" 类型的问题，因为无法评测 RAG 结果
                if question_type == "no answer":
                    continue
                
                # 收集 answer_span 作为 evidence
                evidence = []
                for i in range(1, 6):
                    span = qa_item.get(f"answer_span_{i}", "")
                    if span and span.strip():
                        evidence.append(span.strip())
                
                # 收集 reasoning_steps，也作为 evidence（用于 reasoning 类型问题）
                reasoning_steps = []
                for i in range(1, 6):
                    step = qa_item.get(f"reasoning_step_{i}", "")
                    if step and step.strip():
                        reasoning_steps.append(step.strip())
                        # reasoning_steps 也加入 evidence，用于 recall 计算
                        if step.strip() not in evidence:
                            evidence.append(step.strip())
                
                # 格式化问题
                formatted_question = f'Based on the syllabus "{syllabus_name}", {question}'
                
                qa_pairs.append(StandardQA(
                    question=formatted_question,
                    gold_answers=[answer] if answer else ["Not mentioned"],
                    evidence=evidence,
                    category=question_type,
                    metadata={
                        "id": qa_id,
                        "reasoning_steps": reasoning_steps
                    }
                ))

            # 只有有 QA 对的样本才加入结果
            if qa_pairs:
                standard_samples.append(StandardSample(
                    sample_id=syllabus_name,
                    qa_pairs=qa_pairs
                ))

        return standard_samples

    def _load_from_csv(self, csv_files: List[str]) -> List[StandardSample]:
        """
        从 CSV 文件加载数据。
        
        Args:
            csv_files: CSV 文件路径列表
            
        Returns:
            List[StandardSample]: 标准化样本对象列表
        """
        # 按 syllabus_name 分组
        syllabus_qa_map: Dict[str, List] = {}
        
        for csv_file in csv_files:
            if not os.path.exists(csv_file):
                self.logger.warning(f"CSV file not found: {csv_file}")
                continue
            
            with open(csv_file, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    syllabus_name = row.get('syllabus_name', '')
                    if syllabus_name not in syllabus_qa_map:
                        syllabus_qa_map[syllabus_name] = []
                    syllabus_qa_map[syllabus_name].append(row)
        
        standard_samples = []
        
        for syllabus_name, qa_list in syllabus_qa_map.items():
            qa_pairs = []
            
            for qa_item in qa_list:
                question = qa_item.get("question", "")
                answer = qa_item.get("answer", "")
                question_type = qa_item.get("question_type", "")
                qa_id = qa_item.get("id", "")
                
                # 跳过 "no answer" 类型的问题，因为无法评测 RAG 结果
                if question_type == "no answer":
                    continue
                
                # 收集 answer_span 作为 evidence
                evidence = []
                for i in range(1, 6):
                    span = qa_item.get(f"answer_span_{i}", "")
                    if span and span.strip():
                        evidence.append(span.strip())
                
                # 收集 reasoning_steps，也作为 evidence（用于 reasoning 类型问题）
                reasoning_steps = []
                for i in range(1, 6):
                    step = qa_item.get(f"reasoning_step_{i}", "")
                    if step and step.strip():
                        reasoning_steps.append(step.strip())
                        # reasoning_steps 也加入 evidence，用于 recall 计算
                        if step.strip() not in evidence:
                            evidence.append(step.strip())
                
                # 格式化问题
                formatted_question = f'Based on the syllabus "{syllabus_name}", {question}'
                
                qa_pairs.append(StandardQA(
                    question=formatted_question,
                    gold_answers=[answer] if answer else ["Not mentioned"],
                    evidence=evidence,
                    category=question_type,
                    metadata={
                        "id": qa_id,
                        "reasoning_steps": reasoning_steps
                    }
                ))

            # 只有有 QA 对的样本才加入结果
            if qa_pairs:
                standard_samples.append(StandardSample(
                    sample_id=syllabus_name,
                    qa_pairs=qa_pairs
                ))

        return standard_samples

    def build_prompt(self, qa: StandardQA, context_blocks: List[str]) -> tuple[str, Dict[str, Any]]:
        """
        构建发送给 LLM 的完整提示词。
        
        提示词结构：
        1. 上下文内容（检索到的文档片段）
        2. 无法回答的规则说明
        3. 问题模板
        
        Args:
            qa: 标准化问答对象
            context_blocks: 检索到的上下文文本块列表
            
        Returns:
            tuple[str, Dict[str, Any]]: 
                - 完整的提示词字符串
                - 元数据字典，包含 id
        """
        eff_q = qa.question
        tmpl = QA_PROMPT

        context_text = "\n\n".join(context_blocks)
        full_prompt = f"{context_text}\n\n{MISSING_RULE}\n\n{tmpl.format(eff_q)}"

        meta = {"id": qa.metadata.get("id", "")}
        return full_prompt, meta

    def post_process_answer(self, qa: StandardQA, raw_answer: str, meta: Dict[str, Any]) -> str:
        """
        后处理 LLM 生成的原始答案。
        
        当前实现仅去除首尾空白字符。
        
        Args:
            qa: 标准化问答对象
            raw_answer: LLM 生成的原始答案
            meta: 元数据字典
            
        Returns:
            str: 处理后的答案
        """
        return raw_answer.strip()
